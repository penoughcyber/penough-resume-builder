"""
Resume data extraction from various file formats (PDF, Word, Image).
Modular architecture with fallback chains for robust extraction.
"""

import json
import io
import os
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Any, Tuple, Optional

import pdfplumber
from docx import Document
from PIL import Image
import pytesseract


class BaseExtractor(ABC):
    """Base class for all extractors."""
    
    def __init__(self, file_content: bytes, filename: str):
        self.file_content = file_content
        self.filename = filename
        self.raw_text = ""
        self.extracted_data = self._get_empty_resume()
    
    @staticmethod
    def _get_empty_resume() -> Dict[str, Any]:
        """Returns empty resume structure matching frontend schema."""
        return {
            "name": "",
            "email": "",
            "phone": "",
            "address": "",
            "website": "",
            "websiteLabel": "",
            "linkedin": "",
            "linkedinLabel": "",
            "github": "",
            "githubLabel": "",
            "googleScholar": "",
            "googleScholarLabel": "",
            "portfolio": "",
            "portfolioLabel": "",
            "tagline": "",
            "objective": "",
            "interests": "",
            "competencies": "",
            "technologies": "",
            "experiences": [],
            "education": [],
            "highlights": [],
            "projects": [],
            "books": [],
            "certifications": [],
            "volunteering": [],
            "sectionTitles": {
                "objective": "Professional Profile",
                "experiences": "Professional Experience",
                "education": "Education",
                "highlights": "Key Achievements",
                "interests": "Technical Proficiencies",
                "competencies": "Operational Expertise",
                "projects": "Projects",
                "certifications": "Certifications",
                "volunteering": "Volunteering",
                "books": "Publications",
            },
            "updatedLabel": "",
            "aiResumeBody": ""
        }
    
    @abstractmethod
    def extract(self) -> Tuple[Dict[str, Any], str]:
        """
        Extract resume data from file.
        Returns: (extracted_data dict, raw_text str)
        """
        pass


class PDFExtractor(BaseExtractor):
    """Extract text and structure from PDF files."""
    
    def extract(self) -> Tuple[Dict[str, Any], str]:
        """Extract resume data from PDF using pdfplumber."""
        try:
            with pdfplumber.open(io.BytesIO(self.file_content)) as pdf:
                # Extract all text from all pages
                full_text = ""
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
                
                self.raw_text = full_text
                self._parse_text(full_text)
                
        except Exception as e:
            self.raw_text = f"[PDF extraction error: {str(e)}]"
        
        return self.extracted_data, self.raw_text
    
    def _parse_text(self, text: str):
        """Parse extracted text into resume structure."""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Basic parsing - extract key patterns
        # This is a helper; AI will refine during validation
        current_section = None
        
        for line in lines:
            line_lower = line.lower()
            
            # Contact info patterns
            if '@' in line and '.' in line:
                self.extracted_data['email'] = line.strip()
            elif line_lower.startswith(('phone:', '+')) and any(c.isdigit() for c in line):
                self.extracted_data['phone'] = line.replace('Phone:', '').strip()
            elif line_lower.startswith(('linkedin', 'github', 'website')):
                # Extract URLs
                if 'linkedin' in line_lower:
                    self.extracted_data['linkedin'] = line.split('/')[-1].strip()
                elif 'github' in line_lower:
                    self.extracted_data['github'] = line.split('/')[-1].strip()
            
            # Section headers
            if line_lower in ('professional experience', 'experience', 'work experience'):
                current_section = 'experiences'
            elif line_lower in ('education', 'academic background'):
                current_section = 'education'
            elif line_lower in ('skills', 'technical skills', 'proficiencies'):
                current_section = 'interests'
            elif line_lower in ('certification', 'certifications'):
                current_section = 'certifications'
            elif line_lower in ('project', 'projects'):
                current_section = 'projects'


class WordExtractor(BaseExtractor):
    """Extract text and structure from DOCX files."""
    
    def extract(self) -> Tuple[Dict[str, Any], str]:
        """Extract resume data from DOCX using python-docx."""
        try:
            doc = Document(io.BytesIO(self.file_content))
            
            # Extract all paragraphs and tables
            full_text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text + "\n"
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        full_text += row_text + "\n"
            
            self.raw_text = full_text
            self._parse_text(full_text)
            
        except Exception as e:
            self.raw_text = f"[DOCX extraction error: {str(e)}]"
        
        return self.extracted_data, self.raw_text
    
    def _parse_text(self, text: str):
        """Parse extracted text into resume structure."""
        # Similar to PDF parsing - identify key patterns
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        for line in lines:
            line_lower = line.lower()
            
            # Contact info extraction
            if '@' in line and '.' in line:
                self.extracted_data['email'] = line.strip()
            elif any(digit in line for digit in '0123456789') and ('phone' in line_lower or '+' in line):
                self.extracted_data['phone'] = line.replace('Phone:', '').replace('phone:', '').strip()


class ImageExtractor(BaseExtractor):
    """Extract text from image files using OCR or Vision API."""
    
    def extract(self) -> Tuple[Dict[str, Any], str]:
        """Extract resume data from image file."""
        try:
            # Try OCR first with pytesseract
            image = Image.open(io.BytesIO(self.file_content))
            
            # Resize if very small for better OCR
            if image.width < 800:
                factor = 800 / image.width
                new_size = (int(image.width * factor), int(image.height * factor))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            # Extract text via OCR
            ocr_text = pytesseract.image_to_string(image)
            
            self.raw_text = ocr_text if ocr_text.strip() else "[OCR produced no text - image may need Vision API]"
            self._parse_text(ocr_text)
            
        except Exception as e:
            self.raw_text = f"[Image OCR error: {str(e)}. Consider using Vision API mode.]"
        
        return self.extracted_data, self.raw_text
    
    def _parse_text(self, text: str):
        """Parse OCR text into resume structure."""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Extract contacts and basic info
        for line in lines:
            line_lower = line.lower()
            
            if '@' in line and '.' in line:
                self.extracted_data['email'] = line.strip()
            elif any(digit in line for digit in '0123456789') and (
                'phone' in line_lower or '+' in line or '(' in line
            ):
                self.extracted_data['phone'] = line.replace('Phone:', '').replace('phone:', '').strip()


def get_extractor(file_content: bytes, filename: str) -> BaseExtractor:
    """
    Factory function to select appropriate extractor based on file type.
    
    Args:
        file_content: Binary file content
        filename: Original filename with extension
        
    Returns:
        Appropriate extractor instance
    """
    file_ext = Path(filename).suffix.lower()
    
    if file_ext == '.pdf':
        return PDFExtractor(file_content, filename)
    elif file_ext in ('.docx', '.doc'):
        return WordExtractor(file_content, filename)
    elif file_ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp'):
        return ImageExtractor(file_content, filename)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")


def validate_resume_data(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Validate and normalize extracted resume data to match schema.
    Fills missing required fields and ensures proper structure.
    """
    template = BaseExtractor._get_empty_resume()
    
    # Validate each field type
    for key, value in data.items():
        if key in template:
            if isinstance(template[key], str) and not isinstance(value, str):
                data[key] = str(value) if value else ""
            elif isinstance(template[key], list) and not isinstance(value, list):
                data[key] = []
            elif isinstance(template[key], dict) and not isinstance(value, dict):
                data[key] = {}
    
    # Fill missing keys with defaults
    for key, default_value in template.items():
        if key not in data:
            data[key] = default_value
    
    return data
